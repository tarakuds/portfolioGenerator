#IMPPORTING REQUIRED PACKAGES
from docx import Document

from docx.shared import Inches

import pyttsx3

#DEFINING A FUNCTION FOR TEXT TO SPEECH
def speak(text):
    pyttsx3.speak(text)

#CALLING IMPORTED INBUILT FUNCTIONS
document = Document()

document.add_picture('profile.jpg', Inches(3.0))

speak('Welcome to the automated CV Bot')

#DECLARING VARIABLES FOR PERSONAL PROFILE
document.add_heading('Personal Details')
name = input('What is your name? ')
speak('Thank you' + name + 'Please also provide your email')
email = input('What is your email address? ')
phone_number = int(input('What is your Phone Number? '))
nationality = input('You are a citizen of which country? ')

document.add_paragraph('Name: ' + name)
document.add_paragraph('Email Address: ' + email)
document.add_paragraph('Phone Number: ' + phone_number)
document.add_paragraph('Nationality: ' + nationality)

#DECLARING VARIABLES FOR ABOUT ME SECTION
document.add_heading('About Me')
document.add_paragraph(input('Please give a brief description of yourself '))

#DECLARING VARIABLES FOR EXPERIENCE(S)
document.add_heading('Experience')
experience= document.add_paragraph()
company = input('which company did you work? ')
start_date = input('From what year did you work at ' + company + ' ')
current_position = input('Are you currently working at ' + company + ' ')


experience.add_run(company + ' ' + '|').bold = True

while True:
    if current_position.lower() =='yes':
        end_date = 'till date'
        break
    else:
        end_date = input('What year did you stop work at ' + company + ' ')
        break
 
experience.add_run(start_date + ' - ' + end_date + '\n').italic = True

job_description = input('Briefly describe your job at ' + company + ' ')
experience.add_run(job_description)


while True:
    more_experience = input('Do you have more experiences? ')
    if more_experience.lower() == 'yes':
        experience= document.add_paragraph()
        company = input('which company did you work? ')
        start_date = input('From what year did you work at ' + company)
        while True:
            current_position = input('Are you currently working at ' + company)

            if current_position.lower() =='yes':
                end_date = 'till date'
                break
            else:
                end_date = input('Wwhat year did you stop work at ' + company)
                break

        experience.add_run(company + ' ').bold = True
        experience.add_run(start_date + ' ' + end_date + '\n').italic = True
        job_description = input('Briefly describe your job at ' + company)
        experience.add_run(job_description)
    else:
        break

#DECLARING VARIABLES FOR SKILLSETS
document.add_heading('Skills')
skills = input('Give a list of your skillsets. ')
skillset = document.add_paragraph(skills)
# skillset.add_run(skills)
skillset.style = 'List Bullet'

while True:
    have_more_skills = input('Do you have more skills to add? ')

    if have_more_skills.lower() == 'yes':
        skills = input('List another of your skillsets. ')
        skillset = document.add_paragraph(skills)
        skillset.style = 'List Bullet'
    else:
        break

print('we are done here')

#DECLARING VARIABLES FOR FOOTER SECTION
document.sections[0].footer.paragraphs[0].text = 'Auto generated style by TARA KUDS'

#DECLARING HOW FILE IS TO BE SAVED
document.save('portfolio.docx')